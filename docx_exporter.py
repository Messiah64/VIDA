import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import logging

logger = logging.getLogger(__name__)

class MarkdownToDocxConverter:
    """Convert markdown content to DOCX format"""
    
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
    
    def setup_styles(self):
        """Setup custom styles for the document"""
        styles = self.doc.styles
        
        # Title style
        if 'Custom Title' not in [s.name for s in styles]:
            title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(24)
            title_font.bold = True
            title_font.color.rgb = None  # Default color
            
        # Heading 2 style
        if 'Custom Heading 2' not in [s.name for s in styles]:
            h2_style = styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
            h2_font = h2_style.font
            h2_font.name = 'Calibri'
            h2_font.size = Pt(18)
            h2_font.bold = True
            
        # Heading 3 style
        if 'Custom Heading 3' not in [s.name for s in styles]:
            h3_style = styles.add_style('Custom Heading 3', WD_STYLE_TYPE.PARAGRAPH)
            h3_font = h3_style.font
            h3_font.name = 'Calibri'
            h3_font.size = Pt(14)
            h3_font.bold = True
            
        # Body text style
        if 'Custom Body' not in [s.name for s in styles]:
            body_style = styles.add_style('Custom Body', WD_STYLE_TYPE.PARAGRAPH)
            body_font = body_style.font
            body_font.name = 'Calibri'
            body_font.size = Pt(11)
    
    def parse_markdown_line(self, line):
        """Parse a single line of markdown and return type and content"""
        line = line.strip()
        
        if not line:
            return 'empty', ''
        elif line.startswith('### '):
            return 'h3', line[4:]
        elif line.startswith('## '):
            return 'h2', line[3:]
        elif line.startswith('# '):
            return 'h1', line[2:]
        elif line.startswith('- '):
            return 'bullet', line[2:]
        elif re.match(r'^\d+\. ', line):
            return 'numbered', re.sub(r'^\d+\. ', '', line)
        else:
            return 'paragraph', line
    
    def format_text_with_markdown(self, paragraph, text):
        """Apply markdown formatting (bold, italic) to text within a paragraph"""
        # Handle bold text (**text**)
        parts = re.split(r'(\*\*.*?\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                # Italic text
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                # Regular text
                paragraph.add_run(part)
    
    def add_header(self):
        """Add document header"""
        header_section = self.doc.sections[0]
        header = header_section.header
        header_para = header.paragraphs[0]
        header_para.text = "Video Analysis Report"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page numbers to footer
        footer_section = self.doc.sections[0]
        footer = footer_section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.text = "Page "
    
    def convert_markdown_to_docx(self, markdown_content):
        """Convert markdown content to DOCX format"""
        lines = markdown_content.split('\n')
        current_list_items = []
        current_list_type = None
        
        # Add header
        self.add_header()
        
        # Add title
        title = self.doc.add_paragraph()
        title.style = 'Custom Title'
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.add_run('Video Analysis Report')
        
        # Add spacing
        self.doc.add_paragraph()
        
        for line in lines:
            line_type, content = self.parse_markdown_line(line)
            
            # Handle list items
            if line_type in ['bullet', 'numbered']:
                if current_list_type != line_type:
                    # New list type, finish previous list if exists
                    if current_list_items:
                        self.add_list_to_document(current_list_items, current_list_type)
                        current_list_items = []
                    current_list_type = line_type
                
                current_list_items.append(content)
                continue
            else:
                # Not a list item, finish current list if exists
                if current_list_items:
                    self.add_list_to_document(current_list_items, current_list_type)
                    current_list_items = []
                    current_list_type = None
            
            # Handle other content types
            if line_type == 'empty':
                continue
            elif line_type == 'h1':
                heading = self.doc.add_paragraph()
                heading.style = 'Custom Title'
                self.format_text_with_markdown(heading, content)
            elif line_type == 'h2':
                heading = self.doc.add_paragraph()
                heading.style = 'Custom Heading 2'
                self.format_text_with_markdown(heading, content)
            elif line_type == 'h3':
                heading = self.doc.add_paragraph()
                heading.style = 'Custom Heading 3'
                self.format_text_with_markdown(heading, content)
            elif line_type == 'paragraph':
                para = self.doc.add_paragraph()
                para.style = 'Custom Body'
                self.format_text_with_markdown(para, content)
        
        # Handle remaining list items
        if current_list_items:
            self.add_list_to_document(current_list_items, current_list_type)
    
    def add_list_to_document(self, items, list_type):
        """Add a list to the document"""
        for item in items:
            para = self.doc.add_paragraph()
            para.style = 'List Bullet' if list_type == 'bullet' else 'List Number'
            self.format_text_with_markdown(para, item)

def export_to_docx(markdown_content, output_path):
    """
    Export markdown content to DOCX file
    
    Args:
        markdown_content (str): Markdown content to convert
        output_path (str): Path where DOCX file should be saved
    """
    try:
        logger.info(f"Converting markdown to DOCX: {output_path}")
        
        converter = MarkdownToDocxConverter()
        converter.convert_markdown_to_docx(markdown_content)
        
        # Save document
        converter.doc.save(output_path)
        
        logger.info(f"DOCX export completed: {output_path}")
        
    except Exception as e:
        logger.error(f"Failed to export to DOCX: {str(e)}")
        raise

def create_sample_docx():
    """Create a sample DOCX file for testing"""
    sample_markdown = """# Video Analysis Report

## Executive Summary

This is a sample video analysis report generated from markdown content.

### Key Findings

- **Speaker Identification**: Multiple speakers detected
- **Content Analysis**: Business meeting discussion
- **Sentiment**: Generally positive tone

### Detailed Analysis

The video contains several important segments:

1. **Introduction Phase** (0:00-0:30)
   - Initial greetings and setup
   - Participants settling in

2. **Main Discussion** (0:30-2:00)
   - Core business topics
   - Active participation from all attendees

### Conclusion

The analysis reveals a productive business meeting with clear communication patterns.
"""
    
    export_to_docx(sample_markdown, "sample_report.docx")
    print("Sample DOCX created: sample_report.docx")

if __name__ == "__main__":
    # Create sample DOCX for testing
    create_sample_docx()